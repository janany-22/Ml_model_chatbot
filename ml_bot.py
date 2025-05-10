import streamlit as st
import pandas as pd
import numpy as np
import spacy
import sqlite3
import logging
import json
import os
import pickle
import uuid
import time
from io import StringIO, BytesIO
import difflib
import re
from collections import deque
import pdfplumber
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import docx2txt
from textblob import TextBlob

# Configure logging with both file and stream handlers for visibility
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# File handler for logs
if not os.path.exists("logs"):
    os.makedirs("logs")
file_handler = logging.FileHandler("logs/app.log")
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
logger.addHandler(file_handler)

# Stream handler to display logs in Streamlit
stream_handler = logging.StreamHandler()
stream_handler.setLevel(logging.INFO)
stream_handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
logger.addHandler(stream_handler)

# Ensure data folder exists
DATA_FOLDER = "data"
if not os.path.exists(DATA_FOLDER):
    os.makedirs(DATA_FOLDER)

# Load SpaCy model with error handling
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("Loaded en_core_web_sm model successfully")
except OSError as e:
    logger.warning(f"Failed to load en_core_web_sm: {str(e)}. Using blank English model as fallback.")
    nlp = spacy.blank("en")

# Initialize components
data_store = pd.DataFrame()
available_columns = []  # Store lowercase column names
column_mapping = {}  # Map lowercase to original column names
query_history = deque(maxlen=3)  # Short-term memory for last 3 queries
DATA_METADATA = {}  # Store metadata about the dataset
is_trained = False  # Track training status

# Logging for testing deliverables
query_log = []
output_log = []
accuracy_log = []

def log_query(query, intent, entities):
    entry = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M"),
        "query": query,
        "intent": intent,
        "entities": entities
    }
    query_log.append(entry)
    with open(os.path.join("logs", "query_log.json"), "w") as f:
        json.dump(query_log, f, indent=2)
    logger.info(f"Logged query: {entry}")

def log_output(query, response, logic):
    entry = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M"),
        "query": query,
        "response": response,
        "logic": logic
    }
    output_log.append(entry)
    with open(os.path.join("logs", "output_log.json"), "w") as f:
        json.dump(output_log, f, indent=2)
    logger.info(f"Logged output: {entry}")

def log_accuracy(query, response, is_correct):
    entry = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M"),
        "query": query,
        "response": response,
        "is_correct": is_correct
    }
    accuracy_log.append(entry)
    with open(os.path.join("logs", "accuracy_log.json"), "w") as f:
        json.dump(accuracy_log, f, indent=2)
    logger.info(f"Logged accuracy: {entry}")

# SQLite database setup
DB_NAME = os.path.join(DATA_FOLDER, "chatbot_data.db")

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS user_data (
            id TEXT PRIMARY KEY,
            user_id TEXT,
            upload_time TEXT,
            file_name TEXT,
            data TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# Dynamic Data Type Detection and Metadata Extraction
def identify_data_type(df):
    global DATA_METADATA
    columns = [col.lower() for col in df.columns]
    
    DATA_METADATA = {
        "numeric_columns": [],
        "categorical_columns": [],
        "date_columns": [],
        "data_type": "generic"
    }

    for col in df.columns:
        col_lower = col.lower()
        if pd.api.types.is_numeric_dtype(df[col]):
            DATA_METADATA["numeric_columns"].append(col)
        elif pd.api.types.is_string_dtype(df[col]):
            DATA_METADATA["categorical_columns"].append(col)
        if "date" in col_lower or df[col].astype(str).str.contains(r"\d{4}-\d{2}-\d{2}|\d{2}:\d{2}", na=False, regex=True).any():
            DATA_METADATA["date_columns"].append(col)

    patterns = {
        "school": ["student", "marks", "grade", "class", "subject", "exam", "score"],
        "train": ["train", "departure", "arrival", "station", "platform", "schedule", "classes"],
        "supermarket": ["product", "price", "quantity", "sale", "inventory", "category"],
        "bank": ["transaction", "amount", "credit", "debit", "date", "account"],
        "medical": ["patient", "doctor", "admission", "discharge", "diagnosis", "treatment"],
        "employee": ["employee", "salary", "department", "role", "hire_date"],
        "weather": ["temperature", "humidity", "wind", "precipitation", "forecast"],
        "inventory": ["item", "stock", "quantity", "warehouse"],
        "story": ["chapter", "paragraph", "moral", "story", "tale", "narrative"],
        "shipment": ["order", "status", "delivery", "tracking", "ship", "shipment"]
    }

    scores = {data_type: 0 for data_type in patterns}
    for data_type, keywords in patterns.items():
        for col in columns:
            if any(keyword in col for keyword in keywords):
                scores[data_type] += 1
        if "Extracted Text" in df.columns:
            text_content = " ".join(df["Extracted Text"].dropna().astype(str)).lower()
            for keyword in keywords:
                if keyword in text_content:
                    scores[data_type] += 1

    max_score = max(scores.values())
    if max_score > 0:
        DATA_METADATA["data_type"] = max(k for k, v in scores.items() if v == max_score)

    logger.info(f"Inferred data type: {DATA_METADATA['data_type']}")
    logger.info(f"Data metadata: {DATA_METADATA}")
    return DATA_METADATA["data_type"]

# Data Preprocessing with Enhanced Error Handling
def preprocess_data(file_content, file_name):
    start_time = time.time()
    try:
        file_ext = file_name.split(".")[-1].lower()
        if file_ext == "csv":
            df = pd.read_csv(StringIO(file_content.decode("utf-8", errors="ignore")))
        elif file_ext == "json":
            df = pd.read_json(StringIO(file_content.decode("utf-8", errors="ignore")))
        elif file_ext == "xlsx":
            df = pd.read_excel(BytesIO(file_content))
        elif file_ext in ["txt", "md"]:
            text = file_content.decode("utf-8", errors="ignore").strip()
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            if not lines:
                raise ValueError("No content found in the text file.")
            
            data_dict = {"Extracted Text": lines}
            df = pd.DataFrame(data_dict)
            logger.info(f"Processed {file_ext.upper()} file as unstructured text.")
        elif file_ext == "pdf":
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                all_tables = []
                all_text = []
                table_found = False
                
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        table_found = True
                        for table in tables:
                            if table and table[0]:
                                header = [str(cell) if cell is not None else "Column_" + str(i) for i, cell in enumerate(table[0])]
                                rows = [[str(cell) if cell is not None else "N/A" for cell in row] for row in table[1:]]
                                table_df = pd.DataFrame(rows, columns=header)
                                all_tables.append(table_df)
                    
                    text = page.extract_text()
                    if text:
                        lines = [line.strip() for line in text.split("\n") if line.strip()]
                        all_text.extend(lines)

                if table_found and all_tables:
                    for i, table_df in enumerate(all_tables):
                        table_df.columns = [f"{col}_{i}" if table_df.columns.duplicated()[j] else col 
                                            for j, col in enumerate(table_df.columns)]
                    df = pd.concat(all_tables, ignore_index=True)
                    df = df.dropna(how="all").fillna("N/A")
                    logger.info(f"Extracted table(s) from PDF with columns: {list(df.columns)}")
                    
                    if all_text:
                        table_content = set()
                        for _, row in df.iterrows():
                            for value in row.astype(str):
                                if value and value != "N/A":
                                    table_content.add(value.strip())
                        additional_text = [line for line in all_text if line not in table_content]
                        if additional_text:
                            df["Additional Text"] = " ".join(additional_text)
                            logger.info("Added 'Additional Text' column with non-table content.")
                        else:
                            df["Additional Text"] = "N/A"
                else:
                    if not all_text:
                        raise ValueError("No extractable text or tables found in the PDF.")
                    
                    data_dict = {"Extracted Text": all_text}
                    df = pd.DataFrame(data_dict)
                    logger.info("No tables found in PDF; processed as unstructured text.")
        elif file_ext == "docx":
            try:
                doc = Document(BytesIO(file_content))
                all_tables = []
                all_text = []
                table_found = False

                for table in doc.tables:
                    if table.rows:
                        table_found = True
                        header = [cell.text.strip() if cell.text.strip() else f"Column_{i}" 
                                  for i, cell in enumerate(table.rows[0].cells)]
                        rows = []
                        for row in table.rows[1:]:
                            row_data = [cell.text.strip() if cell.text.strip() else "N/A" for cell in row.cells]
                            rows.append(row_data)
                        table_df = pd.DataFrame(rows, columns=header)
                        all_tables.append(table_df)

                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_text.append(text)

                if table_found and all_tables:
                    for i, table_df in enumerate(all_tables):
                        table_df.columns = [f"{col}_{i}" if table_df.columns.duplicated()[j] else col 
                                            for j, col in enumerate(table_df.columns)]
                    df = pd.concat(all_tables, ignore_index=True)
                    df = df.dropna(how="all").fillna("N/A")
                    logger.info(f"Extracted table(s) from Word document with columns: {list(df.columns)}")
                    
                    if all_text:
                        table_content = set()
                        for _, row in df.iterrows():
                            for value in row.astype(str):
                                if value and value != "N/A":
                                    table_content.add(value.strip())
                        additional_text = [line for line in all_text if line not in table_content]
                        if additional_text:
                            df["Additional Text"] = " ".join(additional_text)
                            logger.info("Added 'Additional Text' column with non-table content from Word document.")
                        else:
                            df["Additional Text"] = "N/A"
                else:
                    if not all_text:
                        raise ValueError("No extractable text or tables found in the Word document.")
                    
                    data_dict = {"Extracted Text": all_text}
                    df = pd.DataFrame(data_dict)
                    logger.info("No tables found in Word document; processed as unstructured text.")
            except (PackageNotFoundError, Exception) as docx_error:
                logger.warning(f"Failed to process .docx with python-docx: {str(docx_error)}. Attempting fallback with docx2txt.")
                try:
                    temp_file_path = os.path.join(DATA_FOLDER, f"temp_{file_name}")
                    with open(temp_file_path, "wb") as temp_file:
                        temp_file.write(file_content)
                    
                    text = docx2txt.process(temp_file_path)
                    os.remove(temp_file_path)
                    
                    if not text.strip():
                        raise ValueError("No extractable text found in the Word document using fallback method.")
                    
                    lines = [line.strip() for line in text.split("\n") if line.strip()]
                    data_dict = {"Extracted Text": lines}
                    df = pd.DataFrame(data_dict)
                    logger.info("Successfully extracted text from Word document using docx2txt fallback.")
                except Exception as fallback_error:
                    logger.error(f"Fallback method failed: {str(fallback_error)}")
                    raise ValueError(
                        f"Unable to process the Word document '{file_name}'. "
                        f"Possible issues: The file may be corrupted, not a valid .docx, or created by an incompatible application. "
                        f"Try opening and re-saving the file in Microsoft Word, or convert it to another supported format (e.g., PDF, TXT)."
                    )
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported types: CSV, JSON, XLSX, PDF, DOCX, TXT, MD.")

        if df.empty:
            raise ValueError("Processed data is empty. Ensure the file contains valid data.")

        df = df.dropna(how="all")
        df = df.fillna(df.select_dtypes(include=np.number).mean().fillna("N/A"))
        for col in df.select_dtypes(include=["object"]).columns:
            df[col] = df[col].str.strip()

        identify_data_type(df)
        
        logger.info(f"Preprocessed {file_name} in {time.time() - start_time:.3f} seconds")
        logger.info(f"DataFrame shape: {df.shape}, Columns: {list(df.columns)}")
        logger.info(f"Data preview (first 5 rows):\n{df.head().to_string()}")
        return df
    except Exception as e:
        logger.error(f"Error preprocessing file {file_name}: {str(e)}")
        st.error(f"Error preprocessing file: {str(e)}")
        return None

# Store Data
def store_data(df, user_id, file_name):
    start_time = time.time()
    data_json = df.to_json()
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    record_id = str(uuid.uuid4())
    c.execute(
        "INSERT INTO user_data (id, user_id, upload_time, file_name, data) VALUES (?, ?, datetime('now'), ?, ?)",
        (record_id, user_id, file_name, data_json)
    )
    conn.commit()
    conn.close()
    df.to_csv(os.path.join(DATA_FOLDER, f"{file_name}_data.csv"), index=False)
    logger.info(f"Stored data for user {user_id}, file {file_name} in {time.time() - start_time:.3f} seconds")

# Load Model from .pkl File
def load_model():
    global data_store, available_columns, column_mapping, DATA_METADATA, is_trained
    model_path = os.path.join(DATA_FOLDER, "data_model.pkl")
    if os.path.exists(model_path):
        try:
            with open(model_path, 'rb') as f:
                model_store = pickle.load(f)
            data_store = model_store["data_store"]
            available_columns = model_store["available_columns"]
            column_mapping = model_store["column_mapping"]
            DATA_METADATA = model_store["metadata"]
            
            # Validate loaded data
            if data_store.empty:
                logger.error("Loaded data_store is empty")
                is_trained = False
                return
            if not available_columns:
                logger.error("No available columns in loaded model")
                is_trained = False
                return
            if not column_mapping:
                logger.error("Column mapping is empty in loaded model")
                is_trained = False
                return
            
            is_trained = True
            logger.info(f"Loaded model from {model_path}")
            logger.info(f"Data store shape: {data_store.shape}")
            logger.info(f"Available columns: {available_columns}")
            logger.info(f"Column mapping: {column_mapping}")
            logger.info(f"Data metadata: {DATA_METADATA}")
            logger.info(f"Data store preview (first 5 rows):\n{data_store.head().to_string()}")
        except Exception as e:
            logger.error(f"Error loading model: {str(e)}")
            is_trained = False
    else:
        logger.info("No saved model found. Training required.")

# Train Model and Save as .pkl
def train_model(df):
    global data_store, available_columns, column_mapping, DATA_METADATA, is_trained
    start_time = time.time()
    try:
        data_store = df.copy()
        column_mapping = {col.lower(): col for col in df.columns}
        available_columns = list(column_mapping.keys())
        logger.info(f"Available columns set to: {available_columns}")
        logger.info(f"Column mapping: {column_mapping}")

        logger.info(f"Training on {len(data_store)} rows and {len(data_store.columns)} columns")

        identifier_col = None
        for col in data_store.select_dtypes(include=["object"]).columns:
            if data_store[col].nunique() == len(data_store) or "id" in col.lower() or "roll" in col.lower():
                identifier_col = col
                break
        if not identifier_col and data_store.select_dtypes(include=["object"]).columns.size > 0:
            identifier_col = data_store.select_dtypes(include=["object"]).columns[0]
        if identifier_col:
            data_store.set_index(identifier_col, inplace=True)
            logger.info(f"Set {identifier_col} as index for row lookup")

        model_store = {
            "data_store": data_store,
            "available_columns": available_columns,
            "column_mapping": column_mapping,
            "metadata": DATA_METADATA
        }
        model_path = os.path.join(DATA_FOLDER, "data_model.pkl")
        with open(model_path, 'wb') as f:
            pickle.dump(model_store, f)
        logger.info(f"Saved model to {model_path}")

        # Validate saved model
        with open(model_path, 'rb') as f:
            loaded_model = pickle.load(f)
        if loaded_model["data_store"].empty:
            logger.error("Saved data_store is empty")
            is_trained = False
            return

        is_trained = True
        logger.info(f"Trained model for {DATA_METADATA['data_type']} data in {time.time() - start_time:.3f} seconds with columns: {available_columns}")
        st.success(f"Data model trained for {DATA_METADATA['data_type']} data in {(time.time() - start_time) * 1000:.3f} ms with columns: {list(column_mapping.values())}")
    except Exception as e:
        logger.error(f"Error training model: {str(e)}")
        st.error(f"Error training model: {str(e)}")
        is_trained = False

# Reset Function to Clear Model and State
def reset_model():
    global data_store, available_columns, column_mapping, DATA_METADATA, is_trained, query_history
    data_store = pd.DataFrame()
    available_columns = []
    column_mapping = {}
    DATA_METADATA = {}
    is_trained = False
    query_history = deque(maxlen=3)
    model_path = os.path.join(DATA_FOLDER, "data_model.pkl")
    if os.path.exists(model_path):
        os.remove(model_path)
        logger.info(f"Deleted saved model at {model_path}")
    if "df" in st.session_state:
        del st.session_state.df
    logger.info("Model and session state reset successfully")
    st.success("Model and session state have been reset. You can now upload a new dataset.")

# Listener
def listener_check(query, user_id):
    start_time = time.time()
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT data FROM user_data WHERE user_id = ?", (user_id,))
    rows = c.fetchall()
    conn.close()

    if not rows:
        logger.warning(f"No data found in database for user {user_id} in {time.time() - start_time:.3f} seconds")
        return None

    for row in rows:
        df = pd.read_json(row[0])
        if "column" in query["entities"] and query["entities"]["column"] in df.columns:
            if not df.equals(data_store):
                logger.info(f"Listener detected new data, returning for retraining in {time.time() - start_time:.3f} seconds")
                return df
    logger.info(f"Listener checked existing data in {time.time() - start_time:.3f} seconds, no retraining needed")
    return None

# Enhanced NLP Processing with SpaCy
def process_query(query):
    global data_store, available_columns, column_mapping, query_history, DATA_METADATA
    start_time = time.time()
    doc = nlp(query.lower())
    intent = "general"
    entities = {}

    # Log the state of key variables for debugging
    logger.info(f"Processing query: {query}")
    logger.info(f"Data store shape: {data_store.shape}")
    logger.info(f"Available columns: {available_columns}")
    logger.info(f"Column mapping: {column_mapping}")
    logger.info(f"Data metadata: {DATA_METADATA}")

    # Token logging
    tokens = [token.text for token in doc]
    logger.info(f"Tokens: {tokens}")

    # Simplified and more flexible intent detection
    intent_keywords = {
        "retrieve": ["what", "find", "get", "is", "tell", "how", "did", "about", "for", "where", "who", "which"],
        "list": ["list", "show", "display", "enumerate", "unique", "what are", "available", "all", "tell me the"],
        "greeting": ["hi", "hello", "good morning", "good afternoon", "good evening"],
        "summary": ["summary", "overview", "describe", "explain", "what is the story"],
        "moral": ["moral", "lesson", "theme", "message", "what did we learn"]
    }

    for intent_type, keywords in intent_keywords.items():
        if any(keyword in query.lower() for keyword in keywords):
            intent = intent_type
            logger.info(f"Detected intent: {intent_type}")
            break

    # Enhanced column and value extraction
    if available_columns:
        phrase = " ".join(tokens).replace("_", " ")
        logger.info(f"Phrase for column matching: {phrase}")

        # Step 1: Direct column matching
        for col_lower in available_columns:
            col_normalized = col_lower.replace("_", " ")
            if col_normalized in phrase or col_lower in phrase:
                entities["column"] = column_mapping[col_lower]
                logger.info(f"Directly matched column '{entities['column']}' from phrase '{phrase}'")
                break

        # Step 2: Fuzzy matching with lower threshold
        if not entities.get("column"):
            for token in tokens:
                matches = difflib.get_close_matches(token, available_columns, n=1, cutoff=0.5)  # Lowered cutoff for broader matching
                if matches:
                    entities["column"] = column_mapping[matches[0]]
                    logger.info(f"Fuzzy matched column '{entities['column']}' for token '{token}'")
                    break

        # Step 3: Fallback to default column based on intent
        if not entities.get("column"):
            if intent == "retrieve":
                entities["column"] = DATA_METADATA.get("numeric_columns", [None])[0] or DATA_METADATA.get("categorical_columns", [None])[0]
            elif intent in ["summary", "moral"]:
                entities["column"] = "Extracted Text" if "Extracted Text" in available_columns else "Additional Text" if "Additional Text" in available_columns else None
            elif intent == "list":
                entities["column"] = DATA_METADATA.get("categorical_columns", [None])[0] or DATA_METADATA.get("numeric_columns", [None])[0]
            if entities.get("column"):
                logger.info(f"Fallback selected column: {entities['column']}")
            else:
                logger.warning("No column could be matched or selected as fallback")

        # Value extraction with improved robustness
        value_candidates = []
        for token in doc:
            if token.pos_ in ["PROPN", "NOUN", "NUM"]:
                value_candidates.append(token.text)
        if value_candidates:
            value = None
            # Look for values after prepositions
            for i, token in enumerate(doc):
                if token.text in ["of", "for", "in", "with"] and i + 1 < len(doc):
                    next_token = doc[i + 1]
                    if next_token.pos_ in ["PROPN", "NOUN", "NUM"]:
                        value = next_token.text
                        break
            if not value:
                # Combine tokens to form potential values
                value = " ".join(value_candidates)
                # Validate against dataset
                for col in data_store.columns:
                    if data_store[col].astype(str).str.contains(value, case=False, na=False).any():
                        break
                else:
                    # Try individual tokens if combined value doesn't match
                    for candidate in value_candidates:
                        for col in data_store.columns:
                            if data_store[col].astype(str).str.contains(candidate, case=False, na=False).any():
                                value = candidate
                                break
                        if value == candidate:
                            break
            entities["value"] = value
            logger.info(f"Matched value '{entities['value']}'")
        else:
            logger.info("No value candidates found in query")

    # Validation and improved user feedback
    if not entities.get("column") and intent not in ["greeting"]:
        message = (f"Could not identify a valid column in your query. Available columns: {', '.join(column_mapping.values()) if column_mapping else 'None'}. "
                   "Try mentioning a column name like 'Physics' or 'Roll No' in your query. For example, 'What is the physics mark of stu0004?'")
        logger.warning(f"No column detected in query: {query}")
        return {"intent": intent, "entities": entities, "raw_query": query, "status": "warning", "message": message}

    if intent == "retrieve" and not entities.get("value"):
        message = (f"Please specify a value (e.g., ID, name, or number) in your query. For example, 'What is the physics mark of stu0004?' "
                   f"Available columns: {', '.join(column_mapping.values()) if column_mapping else 'None'}")
        logger.warning(f"No value detected in query: {query}")
        return {"intent": intent, "entities": entities, "raw_query": query, "status": "warning", "message": message}

    log_query(query, intent, entities)
    query_history.append({"intent": intent, "entities": entities.copy(), "raw_query": query})
    logger.info(f"Processed query: {query}, Intent: {intent}, Entities: {entities} in {time.time() - start_time:.3f} seconds")
    return {"intent": intent, "entities": entities, "raw_query": query}

# Function to Analyze Stories and Extract/Generate Morals
def analyze_story(text):
    doc = nlp(text)
    blob = TextBlob(text)
    
    themes = []
    for chunk in doc.noun_chunks:
        if len(chunk.text.split()) > 1:
            themes.append(chunk.text.lower())
    
    sentiment = blob.sentiment.polarity
    tone = "positive" if sentiment > 0 else "negative" if sentiment < 0 else "neutral"
    
    moral_keywords = ["moral", "lesson", "learned", "taught", "message"]
    sentences = [sent.text for sent in doc.sents]
    explicit_moral = None
    for sentence in sentences[-5:]:
        if any(keyword in sentence.lower() for keyword in moral_keywords):
            explicit_moral = sentence.strip()
            break
    
    if explicit_moral:
        return explicit_moral, "Extracted explicit moral from the story."
    
    common_morals = {
        "friendship": "True friendship brings joy and support.",
        "honesty": "Honesty is the best approach.",
        "hard work": "Hard work leads to success.",
        "kindness": "Kindness always pays off.",
        "greed": "Greed can lead to downfall.",
        "bravery": "Bravery helps overcome challenges."
    }
    
    generated_moral = "Always strive to do your best."
    for theme in themes:
        for key, moral in common_morals.items():
            if key in theme:
                generated_moral = moral
                break
    
    if tone == "negative":
        generated_moral = f"Avoid mistakes: {generated_moral.lower()}"
    
    return generated_moral, f"Generated moral based on themes ({', '.join(themes[:3] if themes else [])}) and sentiment (tone: {tone})."

# Enhanced Query Model Function
def query_model(query, user_id):
    global data_store, available_columns, column_mapping, query_history, DATA_METADATA
    start_time = time.time()
    intent = query["intent"]
    entities = query["entities"]

    # Log the state for debugging
    logger.info(f"Querying model with intent: {intent}, entities: {entities}")
    logger.info(f"Data store shape: {data_store.shape}")
    logger.info(f"Available columns: {available_columns}")
    logger.info(f"Data store preview (first 5 rows):\n{data_store.head().to_string()}")

    # Check if data_store is empty
    if data_store.empty:
        logger.error("Data store is empty")
        return {"status": "error", "message": "No data available. Please upload a file and train the model."}

    try:
        # Check for new data via listener
        new_data = listener_check(query, user_id)
        if new_data is not None:
            train_model(new_data)
            logger.info(f"Retrained model with new data in {time.time() - start_time:.3f} seconds")

        # Function to search text in columns
        def search_text(column, keyword):
            if column not in data_store.columns:
                return []
            results = []
            for text in data_store[column].dropna():
                if isinstance(text, str) and keyword.lower() in text.lower():
                    doc = nlp(text)
                    for sent in doc.sents:
                        if keyword.lower() in sent.text.lower():
                            results.append(sent.text.strip())
            return list(set(results))

        # Handle text-based queries
        if "Extracted Text" in data_store.columns or "Additional Text" in data_store.columns:
            text_columns = [col for col in ["Extracted Text", "Additional Text"] if col in data_store.columns]
            full_text = " ".join(data_store[text_columns[0]].dropna().astype(str)) if text_columns else ""

            if intent == "moral" and full_text:
                moral, logic = analyze_story(full_text)
                response = f"The moral of the story is: {moral}"
                log_output(query["raw_query"], response, logic)
                logger.info(f"Returning moral response: {response}")
                return {"status": "success", "answer": response}

            if intent == "summary" and full_text:
                doc = nlp(full_text)
                sentences = [sent.text for sent in doc.sents]
                summary_length = min(3, len(sentences))
                summary = " ".join(sentences[:summary_length])
                noun_phrases = [chunk.text for chunk in doc.noun_chunks if len(chunk.text.split()) > 1]
                if noun_phrases:
                    summary += f"\nKey themes: {', '.join(noun_phrases[:3])}."
                logic = "Generated summary by extracting initial sentences and key themes."
                response = f"Summary: {summary}"
                log_output(query["raw_query"], response, logic)
                logger.info(f"Returning summary response: {response}")
                return {"status": "success", "answer": response}

            if intent == "retrieve" and entities.get("value"):
                keyword = entities.get("value")
                results = []
                for col in text_columns:
                    results.extend(search_text(col, keyword))
                if results:
                    logic = f"Searched for '{keyword}' in text content."
                    response = f"Information about '{keyword}': {'; '.join(results)}"
                    log_output(query["raw_query"], response, logic)
                    logger.info(f"Returning text search response: {response}")
                    return {"status": "success", "answer": response}
                else:
                    logger.info(f"No text results found for keyword: {keyword}")
                    return {"status": "error", "message": f"No information found for '{keyword}' in the document content. Try a different keyword."}

            if intent == "list" and entities.get("column") in text_columns:
                unique_values = set()
                for col in text_columns:
                    for value in data_store[col].dropna():
                        if isinstance(value, str):
                            unique_values.add(value.strip())
                if unique_values:
                    logic = "Listing unique text lines from document content."
                    response = f"Available content: {', '.join(sorted(unique_values)[:10])}"
                    log_output(query["raw_query"], response, logic)
                    logger.info(f"Returning list response: {response}")
                    return {"status": "success", "answer": response}
                else:
                    logger.info("No unique text values found for listing")
                    return {"status": "error", "message": "No content found in the document."}

        # Handle structured data queries
        if intent == "retrieve":
            col = entities.get("column")
            value = entities.get("value")
            if col and value:
                matching_rows = pd.DataFrame()
                # Check index first
                if data_store.index.name:
                    idx_upper = data_store.index.str.upper()
                    if value.upper() in idx_upper.values:
                        matching_rows = data_store[idx_upper == value.upper()]
                        logger.info(f"Matched value '{value}' in index")
                # If no index match, check other columns
                if matching_rows.empty:
                    for col_name in data_store.columns:
                        if col_name in data_store.select_dtypes(include=["object"]).columns:
                            if data_store[col_name].astype(str).str.upper().str.contains(value.upper(), na=False).any():
                                matching_rows = data_store[data_store[col_name].astype(str).str.upper().str.contains(value.upper(), na=False)]
                                logger.info(f"Matched value '{value}' in column '{col_name}'")
                                break
                        # Check for exact matches in numeric columns
                        elif col_name in data_store.select_dtypes(include=np.number).columns:
                            try:
                                numeric_value = float(value)
                                if (data_store[col_name] == numeric_value).any():
                                    matching_rows = data_store[data_store[col_name] == numeric_value]
                                    logger.info(f"Matched numeric value '{value}' in column '{col_name}'")
                                    break
                            except ValueError:
                                continue
                if not matching_rows.empty:
                    result = matching_rows[col].iloc[0] if col in matching_rows.columns else "N/A"
                    logic = f"Retrieval on column {col} for value {value}"
                    response = f"{col} for {value}: {result}"
                    log_output(query["raw_query"], response, logic)
                    logger.info(f"Returning retrieval response: {response}")
                    return {"status": "success", "answer": response}
                else:
                    logger.info(f"No data found for retrieval query: {value} in {col}")
                    return {"status": "error", "message": f"No data found for {value} in {col}. Ensure the value exists in the dataset or try a different value (e.g., 'stu0004' for Roll No)."}
            else:
                logger.warning("Missing column or value for retrieval query")
                return {"status": "error", "message": "Please specify a column and value for the query. For example, 'What is the physics mark of stu0004?'"}

        elif intent == "list":
            col = entities.get("column")
            if not col:
                logger.warning("No column specified for listing")
                return {"status": "error", "message": "No column specified for listing. Try asking 'List all Roll No' or specify a column name like 'Physics'."}
            if col not in data_store.columns:
                logger.warning(f"Column {col} not found in dataset")
                return {"status": "error", "message": f"Column {col} not found in dataset. Available columns: {', '.join(data_store.columns)}. Try a different column."}
            
            unique_values = set()
            for value in data_store[col].dropna():
                values = [v.strip() for v in str(value).split(",")]
                unique_values.update(values)
            if unique_values:
                logic = f"Listing unique values in column {col}"
                response = f"Available {col}: {', '.join(sorted(unique_values)[:10])}"
                log_output(query["raw_query"], response, logic)
                logger.info(f"Returning list response: {response}")
                return {"status": "success", "answer": response}
            else:
                logger.info(f"No unique values found for column: {col}")
                return {"status": "error", "message": f"No unique values found for {col}. Try a different column."}

        elif intent == "greeting":
            response = "Hello! How can I assist you today?"
            log_output(query["raw_query"], response, "Greeting response")
            logger.info(f"Returning greeting response: {response}")
            return {"status": "success", "answer": response}

        logger.warning("No matching intent handler found")
        return {"status": "error", "message": f"No relevant data found for the query. Try asking about specific columns or values, e.g., 'What is the physics mark of stu0004?' or 'List all Roll No'."}

    except Exception as e:
        logger.error(f"Error querying model: {str(e)}")
        return {"status": "error", "message": f"Error processing query: {str(e)}. Please try rephrasing your query or resetting the model."}

# Streamlit UI
st.title("Intelligent Question-Answering Bot with Online Learning")

st.info("Note: The bot supports text and table extraction from various formats. Scanned or image-based documents may require OCR preprocessing.")
uploaded_file = st.file_uploader("Upload a dataset (CSV, JSON, XLSX, PDF, DOCX, TXT, MD)", type=["csv", "json", "xlsx", "pdf", "docx", "txt", "md"])
user_id = st.text_input("Enter User ID", value="test_user")

if "df" not in st.session_state:
    st.session_state.df = None

# Load model if it exists
if not is_trained:
    load_model()

if uploaded_file is not None:
    start_time = time.time()
    file_content = uploaded_file.getvalue()
    df = preprocess_data(file_content, uploaded_file.name)
    
    if df is not None:
        store_data(df, user_id, uploaded_file.name)
        st.session_state.df = df
        
        if "Extracted Text" in df.columns:
            st.write(f"Extracted {len(df)} lines of text from the {uploaded_file.name.split('.')[-1].upper()}.")
        elif "Additional Text" in df.columns:
            st.write(f"Extracted a table with columns: {', '.join([col for col in df.columns if col != 'Additional Text'])}")
            st.write(f"Additional text outside the table: {df['Additional Text'].iloc[0] if df['Additional Text'].iloc[0] != 'N/A' else 'None'}")
        else:
            st.write(f"Extracted a structured dataset with columns: {', '.join(df.columns)}")
        
        st.write("Preview of the data (first 5 rows):")
        st.write(df.head())
        st.write(f"Dataset uploaded and stored in {(time.time() - start_time) * 1000:.3f} ms!")
    else:
        st.session_state.df = None

if st.session_state.df is not None and not is_trained:
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Train"):
            train_model(st.session_state.df)
    with col2:
        if st.button("Reset"):
            reset_model()
else:
    if st.session_state.df is None and not is_trained:
        st.warning("Please upload a dataset before training.")
    elif is_trained:
        st.info("Model is already trained. You can start asking questions.")
        if st.button("Reset"):
            reset_model()

if is_trained:
    st.subheader("Ask Questions")
    st.write("Example questions: 'What is the physics mark of stu0004?', 'What is the moral of the story?', 'Summarize the document', 'List all Roll No'")
    query = st.text_input("Ask any question about the data")
    if st.button("Submit Query"):
        if not query:
            st.warning("Please enter a question to proceed.")
        elif data_store.empty:
            st.error("No data available. Please upload a file and train the model.")
        else:
            start_time = time.time()
            processed_query = process_query(query)
            processed_query["user_id"] = user_id
            response = query_model(processed_query, user_id)
            
            # Display response
            if response.get("status") == "success":
                st.success(response["answer"])
                is_correct = st.radio(f"Was the answer '{response['answer']}' correct for query '{query}'?", ["Yes", "No"])
                if is_correct:
                    log_accuracy(query, response["answer"], is_correct == "Yes")
            elif response.get("status") == "warning":
                st.warning(response["message"])
            else:
                st.error(response["message"])
            
            logger.info(f"Query processed in {(time.time() - start_time) * 1000:.3f} ms")

    # Display logs in UI
    st.subheader("Logs")
    if os.path.exists("logs/app.log"):
        with open("logs/app.log", "r") as f:
            logs = f.readlines()
        st.text_area("Application Logs", value="".join(logs[-20:]), height=200)
    else:
        st.info("No logs available yet.")
else:
    st.info("Please train the model before asking questions.")